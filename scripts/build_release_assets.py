#!/usr/bin/env python3
"""Build publication assets from text sources for the versioned release."""
from __future__ import annotations

import csv
import hashlib
import re
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer

ROOT = Path(__file__).resolve().parents[1]
MANUSCRIPT_DIR = ROOT / "publication" / "manuscript"
DOCX_OUT = ROOT / "publication" / "Governance_First_Cleanroom_RCM_Working_Paper.docx"
PDF_OUT = ROOT / "publication" / "Governance_First_Cleanroom_RCM_Working_Paper.pdf"
XLSX_OUT = ROOT / "docs" / "Governance_First_Cleanroom_RCM_Controls_and_Data_Dictionary.xlsx"
PNG_OUT = ROOT / "docs" / "architecture.png"


def iter_blocks(text: str):
    for raw in text.splitlines():
        line = raw.rstrip()
        if not line:
            yield ("blank", "")
        elif line.startswith("### "):
            yield ("h3", line[4:])
        elif line.startswith("## "):
            yield ("h2", line[3:])
        elif line.startswith("# "):
            yield ("h1", line[2:])
        elif line.startswith("- "):
            yield ("bullet", line[2:])
        elif re.match(r"^\d+\.\s", line):
            yield ("number", re.sub(r"^\d+\.\s", "", line))
        elif line.startswith("```") or line.startswith("|"):
            yield ("code", line)
        else:
            yield ("p", line)


def clean_inline(text: str) -> str:
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    text = text.replace("**", "").replace("`", "")
    return text


def build_docx(text: str) -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    title_done = False
    for kind, content in iter_blocks(text):
        content = clean_inline(content)
        if kind == "blank":
            continue
        if kind == "h1":
            p = doc.add_heading(content, level=1)
            if not title_done:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_done = True
        elif kind == "h2":
            doc.add_heading(content, level=2)
        elif kind == "h3":
            doc.add_heading(content, level=3)
        elif kind == "bullet":
            doc.add_paragraph(content, style="List Bullet")
        elif kind == "number":
            doc.add_paragraph(content, style="List Number")
        elif kind == "code":
            p = doc.add_paragraph(content)
            p.style = doc.styles["No Spacing"]
            for run in p.runs:
                run.font.name = "Consolas"
                run.font.size = Pt(8)
        else:
            doc.add_paragraph(content)
    doc.save(DOCX_OUT)


def build_pdf(text: str) -> None:
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="TitleCentered", parent=styles["Title"], alignment=TA_CENTER, spaceAfter=18))
    styles.add(ParagraphStyle(name="BodySmall", parent=styles["BodyText"], fontSize=9.5, leading=13, spaceAfter=7))
    styles.add(ParagraphStyle(name="CodeSmall", parent=styles["Code"], fontSize=6.5, leading=8, spaceAfter=3))
    story = []
    title_done = False
    for kind, content in iter_blocks(text):
        content = clean_inline(content)
        content = content.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        if kind == "blank":
            story.append(Spacer(1, 0.06 * inch))
        elif kind == "h1":
            style = styles["TitleCentered"] if not title_done else styles["Heading1"]
            story.append(Paragraph(content, style))
            title_done = True
        elif kind == "h2":
            story.append(Paragraph(content, styles["Heading2"]))
        elif kind == "h3":
            story.append(Paragraph(content, styles["Heading3"]))
        elif kind in {"bullet", "number"}:
            prefix = "• " if kind == "bullet" else "– "
            story.append(Paragraph(prefix + content, styles["BodySmall"]))
        elif kind == "code":
            story.append(Paragraph(content, styles["CodeSmall"]))
        else:
            story.append(Paragraph(content, styles["BodySmall"]))
    doc = SimpleDocTemplate(
        str(PDF_OUT), pagesize=LETTER, rightMargin=0.7 * inch, leftMargin=0.7 * inch,
        topMargin=0.65 * inch, bottomMargin=0.65 * inch,
        title="Governance-First Clean-Room Architecture for Healthcare Revenue Cycle Analytics",
        author="Muhammad Aftab Shaukat",
    )
    doc.build(story)


def style_sheet(ws):
    dark = PatternFill("solid", fgColor="17324D")
    mid = PatternFill("solid", fgColor="305A7A")
    for cell in ws[1]:
        cell.fill = mid
        cell.font = Font(color="FFFFFF", bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions
    for col in range(1, ws.max_column + 1):
        max_len = max(len(str(ws.cell(row=r, column=col).value or "")) for r in range(1, ws.max_row + 1))
        ws.column_dimensions[get_column_letter(col)].width = min(max(max_len + 2, 12), 52)
    for row in ws.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)


def add_csv_sheet(wb: Workbook, name: str, path: Path):
    ws = wb.create_sheet(name)
    with path.open(newline="", encoding="utf-8") as handle:
        for row in csv.reader(handle):
            ws.append(row)
    style_sheet(ws)
    return ws


def build_xlsx() -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Executive Summary"
    rows = [
        ["Governance-First Clean-Room RCM Control Workbook", ""],
        ["Author", "Muhammad Aftab Shaukat"],
        ["ORCID", "https://orcid.org/0009-0009-0342-9877"],
        ["Repository", "https://github.com/aashaukatc/governance-first-cleanroom-rcm"],
        ["Research design", "Design-science reference architecture"],
        ["Evaluation data", "Entirely synthetic, non-PHI, and fictional"],
        ["Dirty scenario", "RED"],
        ["Resolved scenario", "GREEN"],
        ["Clean records evaluated", 200],
        ["Automated tests", "2 expected"],
    ]
    for row in rows:
        ws.append(row)
    ws["A1"].fill = PatternFill("solid", fgColor="17324D")
    ws["A1"].font = Font(color="FFFFFF", bold=True, size=15)
    ws.merge_cells("A1:B1")
    ws.column_dimensions["A"].width = 28
    ws.column_dimensions["B"].width = 66
    for row in ws.iter_rows():
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
    add_csv_sheet(wb, "QA Controls", ROOT / "docs" / "qa_control_matrix.csv")
    ws = wb.create_sheet("Mapping Examples")
    ws.append(["Domain", "Raw value", "Canonical value", "Classification", "Status"])
    specs = [
        ("Payer", "payer_mapping.csv", "payer_type"),
        ("Provider", "provider_mapping.csv", "specialty"),
        ("CPT", "cpt_mapping.csv", "cpt_family"),
        ("Denial", "denial_mapping.csv", "denial_category"),
    ]
    for domain, filename, class_col in specs:
        with (ROOT / "data" / "mappings" / filename).open(newline="", encoding="utf-8") as handle:
            for row in csv.DictReader(handle):
                ws.append([domain, row["raw_value"], row["canonical_value"], row[class_col], "APPROVED"])
    style_sheet(ws)
    wb.save(XLSX_OUT)


def build_png() -> None:
    subprocess.run(["dot", "-Tpng", str(ROOT / "docs" / "architecture.dot"), "-o", str(PNG_OUT)], check=True)


def build_checksums() -> None:
    files = [p for p in ROOT.rglob("*") if p.is_file() and ".git" not in p.parts and p.name != "SHA256SUMS.txt"]
    lines = []
    for path in sorted(files):
        digest = hashlib.sha256(path.read_bytes()).hexdigest()
        lines.append(f"{digest}  {path.relative_to(ROOT).as_posix()}")
    (ROOT / "SHA256SUMS.txt").write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    parts = sorted(MANUSCRIPT_DIR.glob("*.md"))
    if not parts:
        raise FileNotFoundError(f"No manuscript sections found in {MANUSCRIPT_DIR}")
    text = "\n\n".join(path.read_text(encoding="utf-8").rstrip() for path in parts) + "\n"
    build_docx(text)
    build_pdf(text)
    build_xlsx()
    build_png()
    build_checksums()
    for path in [DOCX_OUT, PDF_OUT, XLSX_OUT, PNG_OUT, ROOT / "SHA256SUMS.txt"]:
        print(path.relative_to(ROOT), path.stat().st_size)


if __name__ == "__main__":
    main()
